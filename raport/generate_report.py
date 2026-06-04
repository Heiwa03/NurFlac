#!/usr/bin/env python3
"""Generate NurFlac project report (.docx) following the UTM Proiect de an template."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

IMGS = '/home/heiwa/gitrepos/NurFlac/UML_Diagrams'
OUT  = '/home/heiwa/gitrepos/NurFlac/raport/Proiect_de_an_NurFlac.docx'

def img(name):
    return os.path.join(IMGS, name)

doc = Document()

# ─── PAGE SETUP (A4, margins matching template) ──────────────────────────────
sec = doc.sections[0]
sec.page_width    = Cm(21)
sec.page_height   = Cm(29.7)
sec.top_margin    = Cm(2)
sec.bottom_margin = Cm(2)
sec.left_margin   = Cm(2)
sec.right_margin  = Cm(1)

# Default Normal style: Times New Roman 12pt, no spacing, single line
ns = doc.styles['Normal']
ns.font.name = 'Times New Roman'
ns.font.size = Pt(12)
ns.paragraph_format.space_before        = Pt(0)
ns.paragraph_format.space_after         = Pt(0)
ns.paragraph_format.line_spacing_rule   = WD_LINE_SPACING.SINGLE

# ─── HELPERS ─────────────────────────────────────────────────────────────────

def sfont(run, name='Times New Roman', size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold      = bold
    run.italic    = italic

def blank(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

def cpara(text='', size=12, bold=False, before=0, after=0):
    """Centered paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if text:
        r = p.add_run(text)
        sfont(r, size=size, bold=bold)
    return p

def jpara(text, before=0, after=6, indent=True):
    """Justified body paragraph with optional first-line indent."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    if text:
        r = p.add_run(text)
        sfont(r)
    return p

def bullet(text, before=2, after=2):
    """Bullet list item."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent   = Cm(1.25)
    p.paragraph_format.space_before  = Pt(before)
    p.paragraph_format.space_after   = Pt(after)
    r1 = p.add_run('•  ')
    sfont(r1, bold=True)
    r2 = p.add_run(text)
    sfont(r2)
    return p

def numbered(items, before=2, after=2):
    """Numbered list items."""
    for i, text in enumerate(items, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent  = Cm(1.25)
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after  = Pt(after)
        r1 = p.add_run(f'{i}.  ')
        sfont(r1, bold=True)
        r2 = p.add_run(text)
        sfont(r2)

def h1(text):
    """Chapter title: centered, bold, 14pt, all-caps appearance."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(text)
    sfont(r, size=14, bold=True)
    return p

def h2(text):
    """Section title: left, bold 12pt."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    sfont(r, size=12, bold=True)
    return p

def h3(text):
    """Sub-section: left, bold italic 12pt."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    sfont(r, size=12, bold=True, italic=True)
    return p

def code_block(text):
    """Monospaced code block in Consolas 10pt."""
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Consolas'
        r.font.size = Pt(10)

def figure(path, caption, width=Cm(15)):
    """Centered image with italic caption below."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run()
    try:
        r.add_picture(path, width=width)
    except Exception as e:
        r2 = p.add_run(f'[Image not found: {path}]')
        sfont(r2, italic=True)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after  = Pt(10)
    rc = cap.add_run(caption)
    sfont(rc, size=11, italic=True)

def toc_entry(title, page, indent_cm=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)

    # Dot-leader tab stop at 16cm (~9072 twips)
    pPr = p._p.get_or_add_pPr()
    tabs_el = OxmlElement('w:tabs')
    tab_el  = OxmlElement('w:tab')
    tab_el.set(qn('w:val'), 'right')
    tab_el.set(qn('w:pos'), '9072')
    tab_el.set(qn('w:leader'), 'dot')
    tabs_el.append(tab_el)
    pPr.append(tabs_el)

    r1 = p.add_run(title)
    sfont(r1)
    r2 = p.add_run(f'\t{page}')
    sfont(r2)

# ═════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═════════════════════════════════════════════════════════════════════════════
cpara('MINISTERUL EDUCAŢIEI, CULTURII ȘI CERCETĂRII', bold=True)
cpara('Universitatea Tehnică a Moldovei', size=14, bold=True)
cpara('Facultatea Calculatoare Informatică şi Microelectronică', bold=True)
cpara('Departamentul Ingineria Software și Automatică', bold=True)

blank(5)

cpara('Proiect de an', size=22, bold=True, before=12, after=6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Disciplina: ');  sfont(r, size=14, bold=True)
r = p.add_run('Tehnici și Mecanisme de Proiectare a Produselor Program');  sfont(r, size=14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Tema: ');  sfont(r, size=14, bold=True)
r = p.add_run(
    'NurFlac — Sistem pentru validarea fişierelor audio lossless\n'
    'prin intermediul unui bot Telegram'
)
sfont(r, size=14)

blank(5)

for label, value in [
    ('Student:', 'Ciuc Vlada Marian, TI-XXX'),
    ('Coordonator:', 'Cebotari Daria, asis. univ.'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(f'{label}  _______________  {value}')
    sfont(r, size=11)

blank(3)
cpara('Chișinău, 2026', bold=True)

# ═════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
cpara('Cuprins', size=14, bold=True, after=12)

toc_entry('Introducere', 3)
toc_entry('1. Analiza domeniului de studiu', 4)
toc_entry('1.1 Scopul, obiectivele și cerințele sistemului', 4, indent_cm=1)
toc_entry('1.2 Analiza sistemelor deja existente', 5, indent_cm=1)
toc_entry('2. Realizarea sistemului', 6)
toc_entry('2.1 Proiectarea aplicației', 6, indent_cm=1)
toc_entry('2.2 Descrierea tehnologiilor pentru sistem', 19, indent_cm=1)
toc_entry('2.3 Descrierea la nivel de cod pe module', 20, indent_cm=1)
toc_entry('3. Documentarea produsului realizat', 28)
toc_entry('Concluzii', 30)
toc_entry('Bibliografie', 31)
toc_entry('Anexa A', 32)

# ═════════════════════════════════════════════════════════════════════════════
# INTRODUCERE
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('Introducere')

jpara(
    'Odată cu proliferarea platformelor de distribuție audio digitală, problema '
    'autenticității fişierelor audio lossless a devenit din ce în ce mai acută. '
    'Utilizatorii schimbă frecvent fişiere pretins a fi în format lossless (FLAC, WAV, ALAC, AIFF), '
    'dar în realitate acestea pot fi recodificări ale unor surse cu pierderi de calitate (MP3, AAC, Ogg Vorbis). '
    'Astfel de fişiere „upscale” pot induce în eroare ascultătorii şi degradează '
    'calitatea colecțiilor audio.'
)
jpara(
    'Lucrarea de față prezintă NurFlac, un bot Telegram implementat în .NET 9.0, '
    'care automatizează procesul de validare a fişierelor audio lossless. Sistemul analizează '
    'fiecare fişier primit printr-un lanț de validare în trei etape: verificarea extensiei, '
    'verificarea tipului MIME şi analiza spectrală bazată pe Transformata Fourier Rapidă (FFT). '
    'Analiza spectrală detectează absența energiei în benzile de frecvență înaltă, '
    'trăsătură caracteristică fişierelor lossy recodificate.'
)
jpara(
    'Arhitectura NurFlac aplică cu rigoare douăsprezece şabloane de proiectare din catalogul GoF '
    '(Gang of Four): patru creaționale, patru structurale şi patru comportamentale. Această '
    'alegere nu este ornamentală — fiecare şablon rezolvă o problemă concretă de design '
    'şi contribuie la modularitatea, extensibilitatea şi mentenabilitatea sistemului.'
)
jpara(
    'Lucrarea este structurată în trei capitole principale. Primul capitol analizează domeniul '
    'problemei şi cerințele sistemului, comparând NurFlac cu soluții alternative existente. '
    'Al doilea capitol prezintă realizarea sistemului: diagramele UML ale şabloanelor integrate, '
    'tehnologiile utilizate şi descrierea codului pe module. Al treilea capitol documentează produsul '
    'realizat prin descrierea fluxurilor de utilizare şi a sistemului de configurare.'
)
jpara(
    'Implementarea proiectului a implicat provocări tehnice semnificative: integrarea FFmpeg ca proces extern '
    'pentru extragerea eşantioanelor PCM, analiza FFT cu MathNet.Numerics, gestionarea concurentă a '
    'sesiunilor de album cu tipul Lock din C# 13, şi proiectarea unui sistem de moderare bazat pe maşini '
    'de stări. Toate aceste componente sunt legate prin şabloanele de proiectare alese, formând o '
    'arhitectură coerentă şi testabilă.'
)

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 1
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('1. Analiza domeniului de studiu')

jpara(
    'Domeniul de studiu al proiectului NurFlac se situează la intersecția dintre procesarea '
    'semnalelor audio, arhitectura software şi sistemele de comunicare în timp real. Validarea '
    'autenticității unui fişier audio lossless este o problemă netrivială: simpla '
    'verificare a extensiei sau a antetului fişierului este insuficientă, deoarece orice convertor '
    'poate reambala un flux MP3 într-un container FLAC fără a restaura frecvențele înalte pierdute.'
)

h2('1.1 Scopul, obiectivele și cerințele sistemului')

jpara(
    'Scopul proiectului este crearea unui sistem automatizat, accesibil prin interfața Telegram, '
    'care să verifice autenticitatea fişierelor audio lossless şi să gestioneze utilizatorii '
    'care încearcă să încarce materiale neconforme.'
)

jpara('Obiectivele principale ale sistemului sunt:', indent=False)
for obj in [
    'Acceptarea şi validarea fişierelor în formatele FLAC, WAV, ALAC şi AIFF trimise prin Telegram;',
    'Detectarea fişierelor lossy recodificate ca lossless prin analiza distribuției energetice spectrale (FFT);',
    'Gestionarea sesiunilor de upload pentru albume compuse din mai multe fişiere;',
    'Implementarea unui sistem de moderare a utilizatorilor cu strike-uri, timeout şi ban permanent;',
    'Deduplicarea fişierelor prin hash criptografic (MD5 sau SHA-256, configurabil);',
    'Separarea strictă a comenzilor de utilizator față de comenzile administrative.',
]:
    bullet(obj)

jpara(
    'Cerințele funcționale includ: recepționarea şi descărcarea fişierelor audio din Telegram '
    '(Document sau Audio), validarea în trei etape (extensie → MIME → spectral), raportarea rezultatelor '
    'validării direct în chat, gestionarea stărilor de sesiune per utilizator, şi executarea comenzilor '
    'administrative (ban, timeout, unban) de către administratori autenticați.'
)
jpara(
    'Cerințele nefuncționale vizează: performanța (analiza unui fişier în sub 10 secunde), '
    'securitatea (comenzile administrative sunt invizibile utilizatorilor neautorizați), extensibilitatea '
    '(adaugarea de noi formate prin înregistrare în AudioFormatRegistry), şi persistența (datele '
    'utilizatorilor şi ledgerul de hash-uri sunt stocate în SQLite).'
)

h2('1.2 Analiza sistemelor deja existente')

jpara(
    'Există mai multe instrumente care abordează parțial problema validării audio lossless, '
    'fiecare cu limitări semnificative față de abordarea propusă în NurFlac.'
)
jpara(
    'Spek (Spectrogram Analyzer) este o aplicație desktop open-source care vizualizează spectrograma '
    'unui fişier audio. Utilizatorul poate identifica vizual tăieturile de frecvență caracteristice '
    'fişierelor lossy. Dezavantajele majore sunt: necesitatea instalării locale, absența automatizării '
    '(analiza este manuală) şi lipsa oricărei integrări cu platforme de comunicare.'
)
jpara(
    'Auphonic este un serviciu web de procesare audio care include funcții de normalizare şi filtrare, '
    'dar nu se concentrează pe detectarea autenticității lossless. Nu oferă integrare cu Telegram '
    'şi nu are un sistem de moderare a utilizatorilor.'
)
jpara(
    'SoX (Sound eXchange) este un instrument de linie de comandă pentru procesarea audio care poate extrage '
    'informații despre fişiere, inclusiv frecvența de eşantionare şi adâncimea de bit. '
    'Totuşi, SoX nu efectuează analiza spectrală FFT pentru detectarea recodificărilor lossy şi '
    'nu oferă nicio interfață de utilizator.'
)
jpara(
    'Comparativ cu aceste soluții, NurFlac oferă: automatizare completă prin interfața Telegram, '
    'analiză spectrală FFT integrată, sistem de moderare şi deduplicare, arhitectură bazată '
    'pe şabloane de proiectare care facilitează extensia, şi configurare prin fişier JSON cu '
    'suport pentru variabile de mediu.'
)

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 2
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('2. Realizarea sistemului')

jpara(
    'Realizarea sistemului NurFlac a urmărit principiile arhitecturii curate: separarea '
    'responsabilităților, inversarea dependențelor şi testabilitatea. Cele douăsprezece '
    'şabloane de proiectare nu au fost alese arbitrar — fiecare răspunde unei nevoi arhitecturale '
    'concrete identificate în etapa de analiză. Tabelul de mai jos rezumă maparea şabloanelor '
    'la componentele sistemului.'
)

# Summary table
from docx.oxml.shared import OxmlElement as OE
table = doc.add_table(rows=13, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for cell, txt in zip(hdr, ['Şablon', 'Categorie', 'Componentă principală']):
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    sfont(r, bold=True, size=11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

rows_data = [
    ('Singleton',                'Creațional', 'BotConfigurationManager, SpectralAnalysisEngine'),
    ('Factory Method',           'Creațional', 'CommandFactory, StandardCommandFactory'),
    ('Builder',                  'Creațional', 'AlbumReportBuilder, AlbumReport'),
    ('Abstract Factory',         'Creațional', 'IAudioAnalyzerFactory, LosslessAnalyzerFactory'),
    ('Decorator',                'Structural',  'CommandDecorator, ModerationGuardDecorator, AdminGuardDecorator'),
    ('Facade',                   'Structural',  'AudioPipelineFacade'),
    ('Adapter',                  'Structural',  'FfmpegAdapter'),
    ('Proxy',                    'Structural',  'CachingUserRepositoryProxy'),
    ('State',                    'Comportamental', 'AlbumSession, IdleState, AlbumUploadState'),
    ('Chain of Responsibility',  'Comportamental', 'ExtensionValidationHandler, MimeValidationHandler, SpectralValidationHandler'),
    ('Command',                  'Comportamental', 'IBotCommand, CommandDispatcher, concrete commands'),
    ('Strategy',                 'Comportamental', 'IHashStrategy, Md5HashStrategy, Sha256HashStrategy'),
]
for row, (pattern, cat, component) in zip(table.rows[1:], rows_data):
    for cell, txt in zip(row.cells, [pattern, cat, component]):
        p = cell.paragraphs[0]
        r = p.add_run(txt)
        sfont(r, size=10)

doc.add_paragraph()  # spacing after table

# ─── 2.1 Proiectarea aplicatiei ──────────────────────────────────────────────
h2('2.1 Proiectarea aplicației')

jpara(
    'Această secțiune prezintă diagramele UML pentru fiecare dintre cele douăsprezece '
    'şabloane de proiectare integrate în NurFlac. Diagramele respectă notația UML 2.5 '
    'şi au fost generate cu PlantUML.'
)

# — Creational ————————————————————————————————
h3('2.1.1 Singleton — BotConfigurationManager')
jpara(
    'Şablonul Singleton asigură că există o singură instanță a managerului de '
    'configurație în cadrul aplicației. BotConfigurationManager citeşte appsettings.json '
    'şi variabilele de mediu cu prefixul NURFLAC_ o singură dată, la prima accesare, utilizând '
    'mecanismul Lazy<T> cu LazyThreadSafetyMode.ExecutionAndPublication pentru siguranță în medii '
    'multi-fir. SpectralAnalysisEngine utilizează acelaşi mecanism pentru a partaja motorul FFT între '
    'toți analizatorii audio.', indent=False
)
figure(img('Singleton_BotConfigurationManager.png'), 'Figura 1. Singleton — BotConfigurationManager')

h3('2.1.2 Factory Method — CommandFactory')
jpara(
    'Şablonul Factory Method abstractizează crearea obiectelor IBotCommand. Clasa abstractă '
    'CommandFactory declară metoda fabrică CreateCommand(), pe care StandardCommandFactory o '
    'suprascrie pentru a instanția comanda concretă corespunzătoare tokenului primit. Metoda '
    'non-virtuală Create() aplică decoratorul corect (admin sau moderare) înainte de a returna '
    'comanda invokerului.', indent=False
)
figure(img('FactoryMethod_CommandFactory.png'), 'Figura 2. Factory Method — CommandFactory')

h3('2.1.3 Builder — AlbumReportBuilder')
jpara(
    'Şablonul Builder separă construcția unui raport de album de reprezentarea sa. '
    'AlbumReportBuilder acumulează thread-safe (via C# 13 Lock) intrările de succes şi eşec '
    'pe măsură ce fişierele sunt validate, iar metoda Build() construieşte produsul final '
    'AlbumReport. Acesta expune metoda ToMarkdown() care formatează raportul pentru afişare în Telegram.', indent=False
)
figure(img('Builder_AlbumReportBuilder.png'), 'Figura 3. Builder — AlbumReportBuilder')

h3('2.1.4 Abstract Factory — IAudioAnalyzerFactory')
jpara(
    'Şablonul Abstract Factory permite crearea familiilor de analizoare audio fără a specifica '
    'clasele concrete. IAudioAnalyzerFactory declară metode pentru fiecare format lossless suportat. '
    'LosslessAnalyzerFactory creează analizoare pentru validarea fişierelor autentice, iar '
    'LossyDetectorFactory creează analizoare specializate în detectarea recodificărilor lossy.', indent=False
)
figure(img('AbstractFactory_AudioAnalyzerFactory.png'), 'Figura 4. Abstract Factory — IAudioAnalyzerFactory')

# — Structural ————————————————————————————————
h3('2.1.5 Decorator — CommandDecorator')
jpara(
    'Şablonul Decorator injectează gardã de securitate în lanțul de execuție al comenzilor, '
    'fără a modifica clasele concrete. ModerationGuardDecorator blochează utilizatorii '
    'baniți sau în timeout înaintea execuției comenzilor de utilizator. AdminGuardDecorator '
    'verifică identitatea de administrator şi ignoră silențios comenzile provenite de la '
    'non-administratori — fără a trimite niciun răspuns, astfel încât existența '
    'comenzilor administrative rămâne ascunsă. Cele două decoratoare sunt mutual exclusive.', indent=False
)
figure(img('Decorator_CommandGuards.png'), 'Figura 5. Decorator — CommandDecorator')

h3('2.1.6 Facade — AudioPipelineFacade')
jpara(
    'Şablonul Facade expune un punct de intrare unic pentru întregul lanț de validare audio. '
    'AudioPipelineFacade construieşte intern lanțul Chain of Responsibility (Extension → MIME '
    '→ Spectral), conectează handler-ele şi expune o singură metodă publică '
    'ValidateAsync(). Apelantul nu cunoaşte detaliile de cablare sau traversare ale lanțului.', indent=False
)
figure(img('Facade_AudioPipelineFacade.png'), 'Figura 6. Facade — AudioPipelineFacade')

h3('2.1.7 Adapter — FfmpegAdapter')
jpara(
    'Şablonul Adapter adaptează procesul extern FFmpeg la interfața IFfmpegTool, aşteptată '
    'de analizoarele spectrale. FfmpegAdapter lansează FFmpeg ca subprocess, extrage 30 de secunde de '
    'PCM brut (f32le, mono, 44100 Hz) din mijlocul fişierului şi returnează eşantioanele '
    'ca float[]. Astfel, analizoarele spectrale nu depind de detaliile invocării FFmpeg.', indent=False
)
figure(img('Adapter_FfmpegAdapter.png'), 'Figura 7. Adapter — FfmpegAdapter')

h3('2.1.8 Proxy — CachingUserRepositoryProxy')
jpara(
    'Şablonul Proxy interpune un strat de cache în memoria de lucru între UserService şi '
    'SqliteUserRepository. La fiecare mesaj Telegram, sistemul verifică starea de moderare a utilizatorului; '
    'fără cache, aceasta ar implica o interogare SQLite per mesaj. CachingUserRepositoryProxy '
    'returnează înregistrarea din cache dacă este prezentă (TTL 30s), sau deleagă către '
    'RealSubject şi stochează rezultatul. La UpdateAsync(), intrarea din cache este invalidată.', indent=False
)
figure(img('Proxy_CachingUserRepository.png'), 'Figura 8. Proxy — CachingUserRepositoryProxy')

# — Behavioral ————————————————————————————————
h3('2.1.9 State — AlbumSession')
jpara(
    'Şablonul State permite AlbumSession să se comporte diferit în funcție de starea '
    'curentă, fără instrucțiuni condiționale complexe. IdleState procesează '
    'fişierele individuale direct (validare + înregistrare în ledger) şi permite tranzitia la '
    'AlbumUploadState la comanda /album-upload. AlbumUploadState acumulează fişierele în '
    'PendingFiles şi le validează în bloc la /album-done, generând un raport complet via Builder.', indent=False
)
figure(img('State_AlbumSession.png'), 'Figura 9. State — AlbumSession')

h3('2.1.10 Chain of Responsibility — AudioValidationHandler')
jpara(
    'Şablonul Chain of Responsibility implementează validarea audio ca un lanț de handler-e '
    'specializate. ExtensionValidationHandler verifică extensia fişierului şi respinge formatele '
    'necunoscute sau lossy. MimeValidationHandler validă tipul MIME dacă este disponibil. '
    'SpectralValidationHandler descărcă fişierul local, invocă FFmpeg şi analizează '
    'spectrul FFT. Fiecare handler poate respinge fişierul sau íl poate pasa mai departe în lanț.', indent=False
)
figure(img('ChainOfResponsibility_ValidationPipeline.png'), 'Figura 10. Chain of Responsibility — AudioValidationHandler')

h3('2.1.11 Command — IBotCommand')
jpara(
    'Şablonul Command încapsulează fiecare acțiune a botului ca un obiect IBotCommand cu '
    'metoda ExecuteAsync(). CommandDispatcher (Invoker) apelează ExecuteAsync() fără a cunoaşte '
    'ce face comanda concretă. StandardCommandFactory (ConcreteCreator) creează şi decorează '
    'comenzile înainte ca Invokerul să le primească. HelpCommand acceptă IBotConfiguration '
    'pentru a afişa secțiunea de comenzi admin exclusiv administratorilor.', indent=False
)
figure(img('Command_BotCommands.png'), 'Figura 11. Command — IBotCommand')

h3('2.1.12 Strategy — IHashStrategy')
jpara(
    'Şablonul Strategy permite schimbarea algoritmului de hash utilizat pentru deduplicare fără '
    'a modifica LedgerService. IHashStrategy defineşte interfata cu metodele AlgorithmName şi '
    'ComputeAsync(). Md5HashStrategy şi Sha256HashStrategy sunt implementările concrete. Algoritmul '
    'activ este selectat la pornire din configurație (câmpul Ledger:HashStrategy) şi injectat '
    'în LedgerService.', indent=False
)
figure(img('Strategy_HashStrategy.png'), 'Figura 12. Strategy — IHashStrategy')

# ─── 2.2 Technologies ────────────────────────────────────────────────────────
doc.add_page_break()
h2('2.2 Descrierea tehnologiilor pentru sistem')

jpara(
    'Sistemul NurFlac este construit pe un set de tehnologii alese pentru maturitate, '
    'performanță şi integrare nativă cu ecosistemul .NET 9.0.'
)

h3('.NET 9.0 Worker Service')
jpara(
    'Aplicația rulează ca un Worker Service (IHostedService), gestionat de '
    'Microsoft.Extensions.Hosting. Această arhitectură oferă injectare de dependențe, '
    'logging structurat, configurare ierarhică şi ciclul de viață standard al unui serviciu '
    'de fundal. Compatibilitatea cu containerele Docker şi serviciile systemd este nativă.', indent=False
)

h3('Telegram.Bot v22.8.1')
jpara(
    'Biblioteca Telegram.Bot v22.8.1 oferă un client asincron complet pentru Telegram Bot API. '
    'Sistemul utilizează StartReceiving() cu filtrare pe UpdateType.Message pentru recepționarea '
    'mesajelor, ITelegramBotClient.SendMessage() pentru răspunsuri şi GetFile()/DownloadFile() '
    'pentru descărcarea fişierelor audio.', indent=False
)

h3('SQLite via Microsoft.Data.Sqlite')
jpara(
    'Persistența datelor este asigurată prin două baze de date SQLite independente: '
    'nurflac-users.db (starea utilizatorilor: status, strike count, timeout expiry) şi '
    'nurflac-ledger.db (hash-urile fişierelor validate pentru deduplicare). Alegerea SQLite elimină '
    'dependența de un server de baze de date extern şi simplifică deployment-ul.', indent=False
)

h3('MathNet.Numerics')
jpara(
    'Analiza spectrală FFT este implementată cu MathNet.Numerics, o bibliotecă matematică '
    'de înaltă performanță pentru .NET. Metoda Fourier.Forward() aplică transformata '
    'Fourier rapidă unui vector de eşantioane PCM, iar rezultatul este analizat pentru distribuția '
    'energiei în benzile de frecvență înaltă (19–22 kHz).', indent=False
)

h3('FFmpeg')
jpara(
    'FFmpeg este invocat ca proces extern pentru extragerea eşantioanelor PCM din fişierele audio. '
    'Comanda utilizează -ss 30 -t 30 pentru a extrage 30 de secunde din mijlocul fişierului, '
    'convertite în format PCM brut (-f f32le), mono (-ac 1), la 44100 Hz (-ar 44100). '
    'Această abordare suportă toate formatele acceptate de FFmpeg fără a depinde de '
    'biblioteci native în proces.', indent=False
)

h3('Microsoft.Extensions.Caching.Memory')
jpara(
    'IMemoryCache este utilizat de CachingUserRepositoryProxy pentru a menține în cache '
    'înregistrările utilizatorilor pe durata unui TTL de 30 de secunde. Aceasta reduce '
    'semnificativ numărul de interogări SQLite în scenariile cu trafic intens, fără a '
    'risca servirea unor date depăşite pe o perioadă mai lungă.', indent=False
)

h3('C# 13 — Lock şi Primary Constructors')
jpara(
    'Proiectul utilizează funcționalitățile C# 13: constructorii primari (primary '
    'constructors) pentru injectarea de dependențe concisă şi tipul Lock '
    '(System.Threading.Lock) în AlbumReportBuilder şi AlbumSessionManager pentru sincronizare '
    'thread-safe fără overhead-ul obiectului lock tradițional de tip object.', indent=False
)

# ─── 2.3 Code description ────────────────────────────────────────────────────
doc.add_page_break()
h2('2.3 Descrierea la nivel de cod pe module')

jpara(
    'Această secțiune descrie fiecare şablon de proiectare prin prisma codului sursă, '
    'evidențiind metodele principale şi mecanismele de interacțiune. Codul este prezentat '
    'în format Consolas 10pt, conform convenției.'
)

# Singleton
h3('Singleton — BotConfigurationManager')
jpara(
    'BotConfigurationManager implementează şablonul Singleton utilizând Lazy<T> cu publicare '
    'sigură pe fir. Constructorul privat inițializează IConfiguration din appsettings.json şi '
    'variabilele de mediu. Proprietatea statică Instance expune singura instanță a interfeței '
    'IBotConfiguration:', indent=False
)
code_block(
"""private static readonly Lazy<BotConfigurationManager> _instance =
    new(static () => new BotConfigurationManager(),
        LazyThreadSafetyMode.ExecutionAndPublication);

public static IBotConfiguration Instance => _instance.Value;

private BotConfigurationManager()
{
    _config = new ConfigurationBuilder()
        .SetBasePath(AppContext.BaseDirectory)
        .AddJsonFile("appsettings.json", optional: false)
        .AddEnvironmentVariables("NURFLAC_")
        .Build();
}

public bool IsAdmin(long telegramId) =>
    AdminIds.Contains(telegramId);"""
)

# Factory Method
h3('Factory Method — CommandFactory / StandardCommandFactory')
jpara(
    'CommandFactory declară metoda fabrică abstractă CreateCommand() şi logica '
    'non-virtuală Create() care aplică decoratorul potrivit. Comenzile admin primesc '
    'AdminGuardDecorator; comenzile de utilizator primesc ModerationGuardDecorator:', indent=False
)
code_block(
"""public IBotCommand? Create(string token)
{
    var command = CreateCommand(token);
    if (command is null) return null;

    return IsAdminCommand(token)
        ? new AdminGuardDecorator(command, Config, BotClient,
            _loggerFactory.CreateLogger<AdminGuardDecorator>())
        : new ModerationGuardDecorator(command, UserService, BotClient,
            _loggerFactory.CreateLogger<ModerationGuardDecorator>());
}

protected abstract IBotCommand? CreateCommand(string token);
private static bool IsAdminCommand(string t) => t is "ban" or "timeout" or "unban";"""
)
jpara('StandardCommandFactory suprascrie CreateCommand() şi mapează tokenuriie la instanțe concrete:', indent=False)
code_block(
"""protected override IBotCommand? CreateCommand(string token) => token switch
{
    "start"        => new StartCommand(BotClient),
    "help"         => new HelpCommand(BotClient, Config),
    "formats"      => new FormatsCommand(BotClient, _registry),
    "album-upload" => new AlbumUploadCommand(BotClient, _sessions),
    "album-done"   => new AlbumDoneCommand(BotClient, _sessions),
    "ban"          => new BanCommand(BotClient, UserService),
    "timeout"      => new TimeoutCommand(BotClient, UserService),
    "unban"        => new UnbanCommand(BotClient, UserService),
    _              => null
};"""
)

# Builder
h3('Builder — AlbumReportBuilder')
jpara(
    'AlbumReportBuilder acumulează thread-safe intrările de succes şi eşec, iar '
    'Build() construieşte AlbumReport:', indent=False
)
code_block(
"""public sealed class AlbumReportBuilder(long telegramId)
{
    private readonly Lock _lock = new();
    private readonly List<string> _accepted = [];
    private readonly List<(string Name, string Reason)> _rejected = [];

    public AlbumReportBuilder AddSuccess(string fileName) {
        lock (_lock) { _accepted.Add(fileName); }
        return this;
    }
    public AlbumReportBuilder AddFailure(string fileName, string reason) {
        lock (_lock) { _rejected.Add((fileName, reason)); }
        return this;
    }
    public AlbumReport Build() =>
        new(telegramId, [.._accepted], [.._rejected]);
}"""
)

# Abstract Factory
h3('Abstract Factory — IAudioAnalyzerFactory')
jpara(
    'IAudioAnalyzerFactory declară metodele fabrică pentru fiecare format. '
    'Implementarea concretă returnează analizoare specializate pentru categoria sa:', indent=False
)
code_block(
"""public interface IAudioAnalyzerFactory {
    string Category { get; }
    ISpectralAnalyzer CreateFlacAnalyzer();
    ISpectralAnalyzer CreateWavAnalyzer();
    ISpectralAnalyzer CreateAlacAnalyzer();
    ISpectralAnalyzer CreateAiffAnalyzer();
    ISpectralAnalyzer? CreateForExtension(string ext);
}

// LosslessAnalyzerFactory — ConcreteFactory1
public sealed class LosslessAnalyzerFactory(IFfmpegTool ffmpeg) : IAudioAnalyzerFactory
{
    public string Category => "Lossless";
    public ISpectralAnalyzer CreateFlacAnalyzer() => new FlacAnalyzer(ffmpeg);
    // ...
}"""
)

# Decorator
h3('Decorator — ModerationGuardDecorator / AdminGuardDecorator')
jpara(
    'ModerationGuardDecorator verifică starea de moderare a utilizatorului şi ridică '
    'automat timeout-urile expirate:', indent=False
)
code_block(
"""public override async Task ExecuteAsync(Message message, CancellationToken ct = default)
{
    var telegramId = message.From?.Id ?? 0L;
    var user = await _userService.GetOrCreateAsync(telegramId, ct);

    // Auto-lift expired timeouts
    if (user.IsTimedOut() && user.TimeoutUntil <= DateTime.UtcNow) {
        await _userService.UnbanAsync(telegramId, ct);
        user = await _userService.GetOrCreateAsync(telegramId, ct);
    }
    if (user.IsBanned() || user.IsTimedOut()) {
        await _botClient.SendMessage(message.Chat.Id, "Access denied.", ct);
        return;
    }
    await Inner.ExecuteAsync(message, ct);
}"""
)
jpara('AdminGuardDecorator ignoră silențios comenzile provenite de la non-administratori:', indent=False)
code_block(
"""public override async Task ExecuteAsync(Message message, CancellationToken ct = default)
{
    if (!config.IsAdmin(message.From?.Id ?? 0L))
        return;  // silently dropped — no reply sent, command stays hidden

    await Inner.ExecuteAsync(message, ct);
}"""
)

# Facade
h3('Facade — AudioPipelineFacade')
jpara(
    'AudioPipelineFacade cableză intern lanțul de validare şi expune un singur '
    'punct de intrare:', indent=False
)
code_block(
"""public AudioPipelineFacade(AudioFormatRegistry registry,
    IAudioAnalyzerFactory analyzerFactory, ILoggerFactory loggerFactory)
{
    var extension = new ExtensionValidationHandler(registry, ...);
    var mime      = new MimeValidationHandler(registry, ...);
    var spectral  = new SpectralValidationHandler(analyzerFactory, ...);
    extension.SetNext(mime).SetNext(spectral);
    _chain = extension;
}

public async Task<ValidationResult> ValidateAsync(
    AudioFileContext context, CancellationToken ct = default)
    => await _chain.HandleAsync(context, ct);"""
)

# Adapter
h3('Adapter — FfmpegAdapter')
jpara(
    'FfmpegAdapter adaptează FFmpeg la interfața IFfmpegTool prin invocarea sa ca subprocess:', indent=False
)
code_block(
"""public async Task<float[]> ExtractPcmSamplesAsync(string filePath, CancellationToken ct)
{
    var tmp = Path.GetTempFileName();
    var psi = new ProcessStartInfo("ffmpeg",
        $"-ss 30 -t 30 -i \"{filePath}\" -f f32le -ac 1 -ar 44100 \"{tmp}\" -y")
    { RedirectStandardError = true, UseShellExecute = false };

    using var proc = Process.Start(psi)!;
    await proc.WaitForExitAsync(ct);

    var bytes   = await File.ReadAllBytesAsync(tmp, ct);
    var samples = new float[bytes.Length / 4];
    Buffer.BlockCopy(bytes, 0, samples, 0, bytes.Length);
    return samples;
}"""
)

# Proxy
h3('Proxy — CachingUserRepositoryProxy')
jpara(
    'CachingUserRepositoryProxy interceptează apelurile la IUserRepository şi serveşte '
    'datele din cache la hit:', indent=False
)
code_block(
"""private static readonly TimeSpan CacheTtl = TimeSpan.FromSeconds(30);

public async Task<User> GetOrCreateAsync(long telegramId, CancellationToken ct = default)
{
    if (cache.TryGetValue($"user:{telegramId}", out User? cached) && cached is not null)
        return cached;                          // cache hit

    var user = await inner.GetOrCreateAsync(telegramId, ct);
    cache.Set($"user:{telegramId}", user, CacheTtl);
    return user;
}

public async Task UpdateAsync(User user, CancellationToken ct = default)
{
    await inner.UpdateAsync(user, ct);
    cache.Remove($"user:{user.TelegramId}");    // invalidate on write
}"""
)

# State
h3('State — AlbumSession / AlbumUploadState')
jpara(
    'AlbumUploadState finalizează batch-ul la /album-done, validând toate fişierele '
    'acumulate şi utilizând Builder-ul pentru raport:', indent=False
)
code_block(
"""public async Task<string> HandleAlbumDoneCommandAsync(AlbumSession ctx, CancellationToken ct)
{
    var builder = new AlbumReportBuilder(ctx.TelegramId);
    foreach (var file in ctx.PendingFiles)
    {
        var result = await ctx.Pipeline.ValidateAsync(file, ct);
        if (result.IsValid) {
            if (file.LocalFilePath is not null)
                await ctx.Ledger.RecordAsync(file.LocalFilePath, ctx.TelegramId, ct);
            builder.AddSuccess(file.FileName);
        } else {
            builder.AddFailure(file.FileName, result.RejectionReason!);
        }
    }
    ctx.TransitionTo(new IdleState());
    return builder.Build().ToMarkdown();
}"""
)

# Chain of Responsibility
h3('Chain of Responsibility — ExtensionValidationHandler')
jpara(
    'ExtensionValidationHandler respinge fişierele cu extensii necunoscute sau lossy şi '
    'pasează mai departe cele valide:', indent=False
)
code_block(
"""public override async Task<ValidationResult> HandleAsync(
    AudioFileContext ctx, CancellationToken ct)
{
    if (string.IsNullOrEmpty(ctx.Extension))
        return ValidationResult.Fail("No file extension detected.");

    if (_registry.IsLossy(ctx.Extension))
        return ValidationResult.Fail($"Lossy format '{ctx.Extension}' is not accepted.");

    if (!_registry.IsSupported(ctx.Extension))
        return ValidationResult.Fail($"Unsupported format '{ctx.Extension}'.");

    _logger.LogDebug("[EXT] PASS {Ext}", ctx.Extension);
    return await base.HandleAsync(ctx, ct);   // pass to MimeValidationHandler
}"""
)

# Command
h3('Command — CommandDispatcher / HelpCommand')
jpara(
    'CommandDispatcher rezolvă tokenul comenzii, obține IBotCommand decorat din fabrică '
    'şi íl execută:', indent=False
)
code_block(
"""public async Task DispatchAsync(Message message, CancellationToken ct = default)
{
    var token   = message.Text!.Split(' ')[0].TrimStart('/').ToLowerInvariant();
    var command = factory.Create(token);

    if (command is null) {
        await botClient.SendMessage(message.Chat.Id,
            $"Unknown command: /{token}. Use /help for the command list.",
            cancellationToken: ct);
        return;
    }
    await command.ExecuteAsync(message, ct);
}"""
)
jpara(
    'HelpCommand verifică dacă apelantul este administrator şi afişează '
    'secțiunea de comenzi admin doar în acest caz:', indent=False
)
code_block(
"""public async Task ExecuteAsync(Message message, CancellationToken ct = default)
{
    var isAdmin = config.IsAdmin(message.From?.Id ?? 0L);
    var text = isAdmin
        ? "**NurFlac Commands**\\n/start ... /help ... /formats ...\\n" +
          "\\n**Admin Commands**\\n/ban <id> ... /timeout <id> <h> ... /unban <id>"
          "\\n/resetuser <id> ... /clearusers ... /clearledger"
        : "**NurFlac Commands**\\n/start ... /help ... /formats ...";
    await botClient.SendMessage(message.Chat.Id, text, cancellationToken: ct);
}"""
)

# Strategy
h3('Strategy — IHashStrategy / LedgerService')
jpara(
    'LedgerService utilizează strategia de hash injectată pentru a calcula hash-ul '
    'fişierului şi a verifica duplicatele:', indent=False
)
code_block(
"""public interface IHashStrategy {
    string AlgorithmName { get; }
    Task<string> ComputeAsync(Stream stream, CancellationToken ct);
}

// SHA-256 implementation
public sealed class Sha256HashStrategy : IHashStrategy {
    public string AlgorithmName => "SHA256";
    public async Task<string> ComputeAsync(Stream stream, CancellationToken ct) {
        var bytes = await SHA256.HashDataAsync(stream, ct);
        return Convert.ToHexString(bytes).ToLowerInvariant();
    }
}

// Context (LedgerService) uses whatever IHashStrategy was injected
public async Task<bool> IsDuplicateAsync(string filePath, CancellationToken ct) {
    await using var stream = File.OpenRead(filePath);
    var hash = await hashStrategy.ComputeAsync(stream, ct);
    return await repository.ExistsAsync(hash, ct);
}"""
)

h3('UserService — Logica de strike şi moderare')
jpara(
    'UserService implementează escaladarea automată a sancțiunilor pe baza numărului '
    'de strike-uri acumulate:', indent=False
)
code_block(
"""public async Task ApplyStrikeAsync(long telegramId, int score, CancellationToken ct)
{
    var user = await repository.GetOrCreateAsync(telegramId, ct);
    user.StrikeCount += score;

    if (user.StrikeCount >= 5)       // 5+ strikes → ban permanent
    {
        user.Status       = UserStatus.Banned;
        user.TimeoutUntil = null;
    }
    else if (user.StrikeCount == 4)  // 4 strikes → timeout 72h
    {
        user.Status       = UserStatus.TimedOut;
        user.TimeoutUntil = DateTime.UtcNow.AddHours(72);
    }
    else if (user.StrikeCount == 3)  // 3 strikes → timeout 24h
    {
        user.Status       = UserStatus.TimedOut;
        user.TimeoutUntil = DateTime.UtcNow.AddHours(24);
    }
    await repository.UpdateAsync(user, ct);
}"""
)

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 3 — Documentation
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('3. Documentarea produsului realizat')

jpara(
    'Sistemul NurFlac funcționează ca bot Telegram şi interacționează cu utilizatorii '
    'exclusiv prin interfața de mesagerie. Această secțiune descrie fluxurile principale '
    'de utilizare şi comportamentul sistemului în fiecare scenariu.'
)

h2('3.1 Comenzile de utilizator')

jpara('Utilizatorii neadministratori au acces la următoarele comenzi:', indent=False)
for cmd, desc in [
    ('/start', 'Afişează mesajul de bun venit şi instrucțiunile de bază.'),
    ('/help', 'Listează comenzile disponibile. Administratorii văd şi secțiunea de comenzi admin.'),
    ('/formats', 'Afişează lista formatelor audio lossless acceptate (FLAC, WAV, ALAC, AIFF).'),
    ('/album-upload', 'Inițiază o sesiune de upload pentru un album multi-fişier.'),
    ('/album-done', 'Finalizează sesiunea de album, validează toate fişierele acumulate şi returnează raportul.'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(cmd + '  ')
    sfont(r1, name='Consolas', size=10, bold=True)
    r2 = p.add_run('— ' + desc)
    sfont(r2)

h2('3.2 Comenzile administrative')

jpara(
    'Comenzile administrative sunt vizibile exclusiv administratorilor în /help şi sunt '
    'ignorate silențios dacă sunt trimise de non-administratori — fără niciun '
    'răspuns, astfel încât existența lor rămâne ascunsă:', indent=False
)
for cmd, desc in [
    ('/ban <id>', 'Aplică un ban permanent utilizatorului cu Telegram ID-ul specificat.'),
    ('/timeout <id> <ore>', 'Aplică un timeout temporar de durata specificată (ore).'),
    ('/unban <id>', 'Ridică banul sau timeout-ul activ al utilizatorului specificat.'),
    ('/resetuser <id>', 'Şterge strike-urile utilizatorului şi restaurează statusul Active.'),
    ('/clearusers', 'Şterge toate înregistrările de utilizatori din baza de date.'),
    ('/clearledger', 'Şterge ledger-ul de deduplicare, permiţând re-upload-ul fişierelor.'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(cmd + '  ')
    sfont(r1, name='Consolas', size=10, bold=True)
    r2 = p.add_run('— ' + desc)
    sfont(r2)

h2('3.3 Fluxul de validare a unui fişier individual')

jpara(
    'Când un utilizator trimite un fişier audio, UpdateRouter preia fişierul şi '
    'construieşte un AudioFileContext cu extensia, tipul MIME şi metadatele necesare. '
    'AudioPipelineFacade rulează lanțul de validare în următoarele etape:'
)
numbered([
    'ExtensionValidationHandler verifică că extensia este în lista formatelor lossless suportate '
    'şi nu este lossy (MP3, AAC, OGG, etc.).',
    'MimeValidationHandler verifică că tipul MIME (dacă este prezent) corespunde extensiei '
    'fişierului.',
    'SpectralValidationHandler descărcă fişierul local, invocă FfmpegAdapter pentru '
    'extragerea PCM (30s din mijlocul fişierului), şi analizează spectrul prin '
    'SpectralAnalysisEngine (FFT). Dacă energia în banda 19–22 kHz este sub pragul '
    'MinRatio = 0.005 față de energia în banda audibilă, fişierul este '
    'respins ca posibilă recodificare lossy.',
])
jpara(
    'La finalul validării, rezultatul (acceptat / respins cu motiv) este trimis utilizatorului prin '
    'Telegram. Fişierele acceptate sunt înregistrate în ledgerul SQLite cu hash-ul '
    'criptografic pentru deduplicare.'
)

h2('3.4 Fluxul de upload album')

jpara(
    'Sesiunile de album permit validarea în lot a mai multor fişiere aparținând '
    'aceluiaşi album. Utilizatorul inițiază sesiunea cu /album-upload, trimite fişierele '
    'unul câte unul (fiecare este confirmat că a fost adăugat în coadă), şi '
    'finalizează cu /album-done. Sistemul validează toate fişierele din lot şi generează '
    'un raport Markdown cu lista fişierelor acceptate şi a celor respinse cu motivele respective. '
    'Sesiunile sunt gestionate per utilizator în AlbumSessionManager, cu siguranță thread-safe.'
)

h2('3.5 Sistemul de moderare')

jpara(
    'Fiecare fişier respins generează un strike pentru utilizatorul respectiv. '
    'Pragurile de moderare automată sunt:'
)
for threshold, consequence in [
    ('3 strike-uri', 'timeout automat de 24 de ore'),
    ('4 strike-uri', 'timeout automat de 72 de ore'),
    ('5+ strike-uri', 'ban permanent'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(threshold + ': ')
    sfont(r1, bold=True)
    r2 = p.add_run(consequence)
    sfont(r2)

jpara(
    'Administratorii pot aplica manual ban, timeout sau unban prin comenzile dedicate. Timeout-urile '
    'expirate sunt ridicate automat la următoarea interacțiune a utilizatorului cu botul, '
    'fără intervenție manuală, de către ModerationGuardDecorator.'
)

h2('3.6 Configurarea sistemului')

jpara(
    'Configurarea NurFlac se realizează prin fişierul appsettings.json, cu posibilitatea '
    'de suprascrie orice valoare prin variabile de mediu cu prefixul NURFLAC_ (cu dublu underscore '
    'pentru anidare). Parametrii principali sunt:'
)
for param, desc in [
    ('TelegramBot:Token', 'Token-ul botului Telegram obținut de la BotFather.'),
    ('TelegramBot:AdminIds', 'Lista de ID-uri Telegram ale administratorilor (array JSON).'),
    ('UserManagement:SqlitePath', 'Calea către fişierul bazei de date SQLite a utilizatorilor.'),
    ('Ledger:SqlitePath', 'Calea către fişierul bazei de date SQLite a ledgerului de hash-uri.'),
    ('Ledger:HashStrategy', 'Algoritmul de hash: "MD5" sau "SHA256" (implicit SHA256).'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(param + '  ')
    sfont(r1, name='Consolas', size=10, bold=True)
    r2 = p.add_run('— ' + desc)
    sfont(r2)

# ═════════════════════════════════════════════════════════════════════════════
# CONCLUSIONS
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('Concluzii')

jpara(
    'Proiectul NurFlac demonstrează că aplicarea riguroasă a şabloanelor de proiectare '
    'din catalogul GoF nu reprezintă un scop în sine, ci un mijloc eficace de a construi sisteme '
    'software modulare, extensibile şi mentenabile. Cele douăsprezece şabloane integrate '
    'rezolvă fiecare câte o problemă concretă identificată în etapa de analiză '
    'a cerințelor.'
)
jpara(
    'Principalele realizări ale proiectului sunt: implementarea unui motor de analiză spectrală '
    'FFT integrat nativ în .NET prin MathNet.Numerics; un sistem de moderare bazat pe maşini de '
    'stări (State) cu escaladare automată a sancțiunilor; o arhitectură de comenzi extensibilă '
    '(Command + Factory Method + Decorator) care asigură separarea strictă între comenzile de '
    'utilizator şi cele administrative; şi un pipeline de validare audio în trei etape '
    '(Chain of Responsibility + Facade) extensibil fără a modifica codul existent.'
)
jpara(
    'Obstacolele întâmpinate au inclus: gestionarea diferenței de comportament între '
    'Telegram API pentru fişierele trimise ca Document față de cele trimise ca Audio (rezolvată '
    'prin fallback pe tipul MIME pentru extensie); evitarea avertismentelor CS9107 (captarea dublă în '
    'constructorii primari) prin refactorizarea fabricilor la câmpuri protejate explicite; şi '
    'calibrarea pragului MinRatio din motorul spectral pentru a minimiza falsii pozitivi.'
)
jpara(
    'Din perspectiva principiilor SOLID, arhitectura respectă: Single Responsibility (fiecare clasă '
    'are o singură responsabilitate), Open/Closed (noi formate şi strategii se adaugă fără '
    'modificarea codului existent), Liskov Substitution (toate analizoarele şi handler-ele sunt '
    'interschimbabile prin interfețele lor), Interface Segregation (interfețe granulare: '
    'IUserRepository, IHashStrategy, IFfmpegTool) şi Dependency Inversion (dependențele sunt '
    'injectate, nu instanțiate direct).'
)
jpara(
    'Ca evoluții viitoare, sistemul ar putea fi extins cu: suport pentru formate DSD şi MQA, '
    'un dashboard web pentru administrare, notificări push pentru administratori la depăşirea '
    'pragurilor de moderare, şi integrarea unui model de machine learning pentru detecția mai '
    'precisă a recodificărilor.'
)

# ═════════════════════════════════════════════════════════════════════════════
# BIBLIOGRAPHY
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('Bibliografie')

bibliography = [
    'Gamma, E., Helm, R., Johnson, R., Vlissides, J. (1994). Design Patterns: Elements of Reusable '
    'Object-Oriented Software. Addison-Wesley Professional.',
    'Microsoft Corporation. (2024). .NET 9.0 Documentation. [Resursă electronică]. '
    'Disponibil: https://learn.microsoft.com/dotnet/',
    'Telegram. (2024). Telegram Bot API Documentation. [Resursă electronică]. '
    'Disponibil: https://core.telegram.org/bots/api',
    'MathNet Project. (2024). MathNet.Numerics — Numerical Mathematics for .NET. [Resursă '
    'electronică]. Disponibil: https://numerics.mathdotnet.com/',
    'FFmpeg Team. (2024). FFmpeg Documentation. [Resursă electronică]. '
    'Disponibil: https://ffmpeg.org/documentation.html',
    'Martin, R. C. (2008). Clean Code: A Handbook of Agile Software Craftsmanship. Prentice Hall.',
    'Skeet, J. (2019). C# in Depth, 4th Edition. Manning Publications.',
    'Microsoft Corporation. (2024). Microsoft.Data.Sqlite Documentation. [Resursă electronică]. '
    'Disponibil: https://learn.microsoft.com/dotnet/standard/data/sqlite/',
]
for i, ref in enumerate(bibliography, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before     = Pt(2)
    p.paragraph_format.space_after      = Pt(4)
    p.paragraph_format.left_indent      = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    r = p.add_run(f'[{i}] {ref}')
    sfont(r)

# ═════════════════════════════════════════════════════════════════════════════
# ANNEXE A
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('Anexa A — Structura directorului proiectului')

jpara(
    'Mai jos este prezentată structura completă a directorului proiectului NurFlac, '
    'organizată conform principiilor Clean Architecture. Fiecare subdirector corespunde '
    'unui domeniu clar delimitat şi unuia sau mai multor şabloane de proiectare.'
)

code_block(
"""NurFlac/
├── Configuration/          # Singleton   — IBotConfiguration, BotConfigurationManager
├── Audio/
│   ├── Abstractions/       # IFfmpegTool, ISpectralAnalyzer, ISpectralAnalysisEngine
│   ├── Adapters/           # Adapter     — FfmpegAdapter
│   ├── Analyzers/          # BaseSpectralAnalyzer + format-specific analyzers
│   ├── Engine/             # Singleton   — SpectralAnalysisEngine (FFT)
│   ├── Factories/          # Abs.Factory — IAudioAnalyzerFactory, LosslessAnalyzerFactory
│   ├── Facade/             # Facade      — AudioPipelineFacade
│   ├── Models/             # AudioFileContext, AudioFormatRegistry, ValidationResult
│   └── Pipeline/           # Chain of Responsibility — validation handlers
├── Album/
│   ├── Report/             # Builder     — AlbumReportBuilder, AlbumReport
│   ├── States/             # State       — IdleState, AlbumUploadState
│   ├── AlbumSession.cs     # State Context
│   └── AlbumSessionManager.cs
├── Commands/
│   ├── Abstractions/       # IBotCommand
│   ├── Concrete/           # ConcreteCommands: Start, Help, Formats, AlbumUpload, ...
│   ├── Decorators/         # Decorator   — ModerationGuardDecorator, AdminGuardDecorator
│   ├── Factory/            # Factory Method — CommandFactory, StandardCommandFactory
│   └── CommandDispatcher.cs   # Command Invoker
├── Ledger/
│   ├── Hashing/            # Strategy    — IHashStrategy, Md5HashStrategy, Sha256HashStrategy
│   ├── LedgerService.cs    # Strategy Context
│   └── SqliteLedgerRepository.cs
├── Users/
│   ├── Entities/           # User, UserStatus
│   ├── CachingUserRepositoryProxy.cs   # Proxy
│   ├── SqliteUserRepository.cs         # RealSubject
│   └── UserService.cs                  # IUserService implementation
├── Infrastructure/
│   └── Telegram/
│       ├── UpdateRouter.cs              # Message routing & file dispatch
│       └── TelegramBotWorker.cs         # BackgroundService host
├── Storage/                # Organized & flat audio file storage
├── UML_Diagrams/           # 12 PlantUML diagrams + PNG renders
├── appsettings.json        # Configuration (token, admin IDs, DB paths, hash strategy)
└── Program.cs              # DI composition root"""
)

doc.save(OUT)
print(f'Report saved to: {OUT}')
